"""
Streamlit app: Extract text from image/CSV/XLSX/PDF/TXT/paste → LLM → Output → Download
Supports uploading up to 10 files at once.
"""

import io
import json
import os
import streamlit as st
import pandas as pd
import pytesseract
from PIL import Image
from google import genai
import pypdf
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── 1. LLM CALLS ─────────────────────────────────────────────────────────────

def call_llm_csv(text: str, column_names: list) -> str:
    """Ask Gemini to extract structured data as a JSON array for export."""
    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    if column_names:
        column_instruction = (
            f"The output must have exactly these columns in this order: {', '.join(column_names)}. "
            "If a value is not found for a column, use an empty string."
        )
    else:
        column_instruction = "Determine the best column names from the content."
    prompt = f"""
    Extract all structured data from the text below and return it as a JSON array
    of flat objects with the same keys in every object. Return ONLY valid JSON, no explanation.

    {column_instruction}

    Text:
    {text}
    """
    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return response.text


def call_llm_word(text: str) -> str:
    """
    Ask Gemini to extract content into a structured JSON document
    with sections, prose, and tables for a polished Word output.
    """
    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    prompt = f"""
    Read the text below and convert it into a structured document.
    Return a JSON object with the following keys:
    - "title": a short document title (string)
    - "sections": a list of section objects. Each section object must have:
        - "heading": section heading (string)
        - "content": a list of content blocks, where each block is one of:
            - {{"type": "paragraph", "text": "..."}}
            - {{"type": "bullets", "items": ["...", "..."]}}
            - {{"type": "table", "headers": ["col1", "col2"], "rows": [["val", "val"], ...]}}

    Organize the content logically. Use tables where data is structured.
    Use bullet points for lists. Use paragraphs for narrative content.
    Return ONLY valid JSON, no explanation.

    Text:
    {text}
    """
    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return response.text


def call_llm_slide(text: str, slide_specs: list) -> str:
    """Ask Gemini to generate slides based on user-defined specs."""
    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    spec_lines = "\n".join(
        f"Slide {i+1}: Title = \"{s['title']}\", Format = {s['format']}"
        for i, s in enumerate(slide_specs)
    )
    format_instructions = """
Format rules:
- "Title Slide": return keys "title" and "subtitle" (a short one-line summary)
- "Bullet Points": return keys "title" and "bullets" (list of 3-5 strings)
- "Table": return keys "title", "headers" (list of column names), and "rows" (list of lists)
- "Two Column": return keys "title", "left_heading", "left_points" (list), "right_heading", "right_points" (list)
- "Timeline": return keys "title" and "events" (list of objects with "date" and "description")
- "Action Items": return keys "title" and "actions" (list of objects with "task", "owner", "due_date")
"""
    prompt = f"""
    Using the source text below, generate a JSON array of presentation slides.
    Each slide must follow the exact format specified.

    Slide specifications:
    {spec_lines}

    {format_instructions}

    Return ONLY a valid JSON array, one object per slide, in the same order as the specs. No explanation.

    Source text:
    {text}
    """
    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return response.text


def call_llm_summary(text: str) -> str:
    """Ask Gemini to produce a summary and action items as JSON."""
    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    prompt = f"""
    Read the text below and return a JSON object with exactly two keys:
    - "summary": a list of 3-5 strings, each a key takeaway from the text
    - "action_items": a list of objects with keys "task", "owner" (if mentioned, else "Unassigned"), and "due_date" (if mentioned, else "TBD")

    Return ONLY valid JSON, no explanation.

    Text:
    {text}
    """
    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return response.text


# ── 2. TEXT EXTRACTION HELPERS ───────────────────────────────────────────────

def extract_text_from_image(uploaded_file) -> str:
    image = Image.open(uploaded_file)
    return pytesseract.image_to_string(image)

def extract_text_from_csv(uploaded_file) -> str:
    df = pd.read_csv(uploaded_file)
    return df.to_string(index=False)

def extract_text_from_xlsx(uploaded_file) -> str:
    df = pd.read_excel(uploaded_file, engine="openpyxl")
    return df.to_string(index=False)

def extract_text_from_pdf(uploaded_file) -> str:
    reader = pypdf.PdfReader(uploaded_file)
    pages = [page.extract_text() for page in reader.pages if page.extract_text()]
    return "\n\n".join(pages)

def extract_text_from_txt(uploaded_file) -> str:
    return uploaded_file.read().decode("utf-8")

def extract_text_from_file(uploaded_file, input_mode: str) -> str:
    if input_mode == "Image (OCR)":        return extract_text_from_image(uploaded_file)
    elif input_mode == "PDF":              return extract_text_from_pdf(uploaded_file)
    elif input_mode == "CSV":              return extract_text_from_csv(uploaded_file)
    elif input_mode == "Excel (.xlsx)":    return extract_text_from_xlsx(uploaded_file)
    elif input_mode == "Text File (.txt)": return extract_text_from_txt(uploaded_file)
    return ""


# ── 3. RESPONSE PARSERS ──────────────────────────────────────────────────────

def parse_json(response: str):
    """Strip markdown fences and parse JSON from LLM response."""
    cleaned = response.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    return json.loads(cleaned)

def parse_csv_response(response: str) -> pd.DataFrame:
    data = parse_json(response)
    if not isinstance(data, list):
        raise ValueError("Expected a JSON array of objects.")
    return pd.DataFrame(data)

def parse_summary_response(response: str) -> dict:
    data = parse_json(response)
    if "summary" not in data or "action_items" not in data:
        raise ValueError("Response missing 'summary' or 'action_items' keys.")
    return data

def parse_slide_response(response: str) -> list:
    data = parse_json(response)
    if not isinstance(data, list):
        raise ValueError("Expected a JSON array of slide objects.")
    return data

def parse_word_response(response: str) -> dict:
    data = parse_json(response)
    if "title" not in data or "sections" not in data:
        raise ValueError("Response missing 'title' or 'sections' keys.")
    return data


# ── 4. WORD DOCUMENT BUILDERS ─────────────────────────────────────────────────

def add_table_to_doc(doc, headers: list, rows: list):
    """Helper to add a formatted table to a Word doc."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark blue background
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val) if val is not None else ""


def build_word_doc_from_sections(doc_data: dict) -> bytes:
    """Build a polished Word doc from structured section data."""
    doc = Document()

    # Title
    title_para = doc.add_heading(doc_data.get("title", "Document"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section in doc_data.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        for block in section.get("content", []):
            btype = block.get("type", "paragraph")
            if btype == "paragraph":
                doc.add_paragraph(block.get("text", ""))
            elif btype == "bullets":
                for item in block.get("items", []):
                    doc.add_paragraph(item, style="List Bullet")
            elif btype == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                if headers and rows:
                    add_table_to_doc(doc, headers, rows)
                    doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_word_doc_from_slides(slides: list) -> bytes:
    """Convert slide data into a formatted Word document."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for i, slide in enumerate(slides):
        fmt = slide.get("format", "Bullet Points")
        if i > 0:
            doc.add_page_break()

        if fmt == "Title Slide":
            p = doc.add_heading(slide.get("title", ""), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub = doc.add_paragraph(slide.get("subtitle", ""))
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif fmt == "Bullet Points":
            doc.add_heading(slide.get("title", ""), level=1)
            for b in slide.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

        elif fmt == "Table":
            doc.add_heading(slide.get("title", ""), level=1)
            headers = slide.get("headers", [])
            rows = slide.get("rows", [])
            if headers and rows:
                add_table_to_doc(doc, headers, rows)

        elif fmt == "Two Column":
            doc.add_heading(slide.get("title", ""), level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = slide.get("left_heading", "")
            table.rows[0].cells[1].text = slide.get("right_heading", "")
            for cell in table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True
            left_pts  = slide.get("left_points", [])
            right_pts = slide.get("right_points", [])
            for k in range(max(len(left_pts), len(right_pts))):
                row = table.add_row()
                row.cells[0].text = left_pts[k]  if k < len(left_pts)  else ""
                row.cells[1].text = right_pts[k] if k < len(right_pts) else ""

        elif fmt == "Timeline":
            doc.add_heading(slide.get("title", ""), level=1)
            for e in slide.get("events", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{e.get('date', '')} — ").bold = True
                p.add_run(e.get("description", ""))

        elif fmt == "Action Items":
            doc.add_heading(slide.get("title", ""), level=1)
            actions = slide.get("actions", [])
            if actions:
                add_table_to_doc(doc, ["Task", "Owner", "Due Date"],
                                 [[a.get("task",""), a.get("owner",""), a.get("due_date","")] for a in actions])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 5. SLIDE RENDERER ────────────────────────────────────────────────────────

SLIDE_FORMATS = ["Title Slide", "Bullet Points", "Table", "Two Column", "Timeline", "Action Items"]

def render_slides_in_app(slides: list):
    for i, slide in enumerate(slides):
        fmt = slide.get("format", "Bullet Points")
        st.markdown("---")
        st.markdown(f"### Slide {i+1}")
        if fmt == "Title Slide":
            st.markdown(f"# {slide.get('title', '')}")
            st.markdown(f"*{slide.get('subtitle', '')}*")
        elif fmt == "Bullet Points":
            st.markdown(f"**{slide.get('title', '')}**")
            for b in slide.get("bullets", []):
                st.markdown(f"- {b}")
        elif fmt == "Table":
            st.markdown(f"**{slide.get('title', '')}**")
            headers, rows = slide.get("headers", []), slide.get("rows", [])
            if headers and rows:
                st.dataframe(pd.DataFrame(rows, columns=headers), use_container_width=True)
        elif fmt == "Two Column":
            st.markdown(f"**{slide.get('title', '')}**")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**{slide.get('left_heading','Left')}**")
                for p in slide.get("left_points", []): st.markdown(f"- {p}")
            with col2:
                st.markdown(f"**{slide.get('right_heading','Right')}**")
                for p in slide.get("right_points", []): st.markdown(f"- {p}")
        elif fmt == "Timeline":
            st.markdown(f"**{slide.get('title', '')}**")
            for e in slide.get("events", []):
                st.markdown(f"- **{e.get('date','')}** — {e.get('description','')}")
        elif fmt == "Action Items":
            st.markdown(f"**{slide.get('title', '')}**")
            actions = slide.get("actions", [])
            if actions:
                st.dataframe(pd.DataFrame(actions), use_container_width=True)


# ── 6. STREAMLIT UI ──────────────────────────────────────────────────────────

def main():
    st.set_page_config(page_title="Navistone Content Extractor", layout="centered")
    st.title("📋 Navistone Content Extractor")
    st.caption("Upload up to 10 files or paste text — then choose how to extract it.")

    # ── Input ────────────────────────────────────────────────────────────────
    st.header("1 · Provide input")
    input_mode = st.radio(
        "Input type",
        ["Image (OCR)", "PDF", "CSV", "Excel (.xlsx)", "Text File (.txt)", "Paste text"],
        horizontal=True
    )

    raw_text = ""

    if input_mode == "Paste text":
        raw_text = st.text_area("Paste your text here", height=200,
                                placeholder="Paste an email, meeting notes, campaign details…")
    else:
        file_type_map = {
            "Image (OCR)":        ["png", "jpg", "jpeg"],
            "PDF":                ["pdf"],
            "CSV":                ["csv"],
            "Excel (.xlsx)":      ["xlsx"],
            "Text File (.txt)":   ["txt"],
        }
        uploaded_files = st.file_uploader(
            f"Upload up to 10 {input_mode} files",
            type=file_type_map[input_mode],
            accept_multiple_files=True,
        )
        if uploaded_files and len(uploaded_files) > 10:
            st.warning("Maximum 10 files allowed. Only the first 10 will be used.")
            uploaded_files = uploaded_files[:10]

        if uploaded_files:
            all_texts = []
            for f in uploaded_files:
                with st.spinner(f"Reading {f.name}..."):
                    try:
                        all_texts.append(f"--- {f.name} ---\n{extract_text_from_file(f, input_mode)}")
                    except Exception as e:
                        st.error(f"Could not read {f.name}: {e}")
            if all_texts:
                raw_text = "\n\n".join(all_texts)
                st.success(f"✅ {len(all_texts)} file(s) loaded successfully.")
                with st.expander("Preview extracted text", expanded=False):
                    st.text_area("Combined text", raw_text, height=200, disabled=True)

    # ── Extraction mode ──────────────────────────────────────────────────────
    st.header("2 · Choose extraction type")
    extraction_mode = st.radio(
        "What do you want?",
        ["Extract to CSV", "Extract to Excel", "Extract to Word", "Slide Ready", "Summary & Action Items"],
        horizontal=True,
    )

    # ── Sub-options ──────────────────────────────────────────────────────────
    column_names = []
    slide_specs  = []

    if extraction_mode in ("Extract to CSV", "Extract to Excel"):
        st.markdown("**Column options**")
        use_custom_columns = st.toggle("Let me define the columns", value=False)
        if use_custom_columns:
            col1, col2 = st.columns([1, 2])
            with col1:
                num_columns = st.number_input("Number of columns", min_value=1, max_value=20, value=3, step=1)
            with col2:
                columns_input = st.text_input("Column names",
                                              placeholder="e.g. Client Name, Campaign Budget, Start Date")
            if columns_input.strip():
                column_names = [c.strip() for c in columns_input.split(",") if c.strip()]
                while len(column_names) < num_columns:
                    column_names.append(f"Column {len(column_names) + 1}")
                column_names = column_names[:num_columns]
                st.caption(f"✅ Columns: {', '.join(column_names)}")

    elif extraction_mode == "Slide Ready":
        st.markdown("**Define your slides**")
        num_slides = st.number_input("How many slides?", min_value=1, max_value=20, value=3, step=1)
        st.caption("Set a title and format for each slide below.")
        for i in range(int(num_slides)):
            col1, col2 = st.columns([2, 2])
            with col1:
                title = st.text_input(f"Slide {i+1} title",
                                      placeholder="e.g. Overview",
                                      key=f"slide_title_{i}")
            with col2:
                fmt = st.selectbox(f"Slide {i+1} format", SLIDE_FORMATS,
                                   index=0 if i == 0 else 1,
                                   key=f"slide_format_{i}")
            slide_specs.append({"title": title or f"Slide {i+1}", "format": fmt})

    # ── Run ──────────────────────────────────────────────────────────────────
    st.header("3 · Extract")
    run = st.button("🚀 Run", disabled=not raw_text.strip())

    if run:
        with st.spinner("Sending to Gemini..."):
            try:
                if extraction_mode in ("Extract to CSV", "Extract to Excel"):
                    llm_output = call_llm_csv(raw_text, column_names)
                elif extraction_mode == "Extract to Word":
                    llm_output = call_llm_word(raw_text)
                elif extraction_mode == "Slide Ready":
                    llm_output = call_llm_slide(raw_text, slide_specs)
                else:
                    llm_output = call_llm_summary(raw_text)
            except Exception as e:
                st.error(f"Error: {e}")
                st.stop()

        # ── Results ──────────────────────────────────────────────────────────
        st.header("4 · Results")

        if extraction_mode == "Extract to CSV":
            try:
                df = parse_csv_response(llm_output)
                st.dataframe(df, use_container_width=True)
                st.download_button("⬇️ Download as CSV",
                                   data=df.to_csv(index=False).encode("utf-8"),
                                   file_name="extracted_data.csv", mime="text/csv")
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Could not parse response: {e}")
                st.code(llm_output)

        elif extraction_mode == "Extract to Excel":
            try:
                df = parse_csv_response(llm_output)
                st.dataframe(df, use_container_width=True)
                buf = io.BytesIO()
                df.to_excel(buf, index=False, engine="openpyxl")
                st.download_button("⬇️ Download as Excel",
                                   data=buf.getvalue(),
                                   file_name="extracted_data.xlsx",
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Could not parse response: {e}")
                st.code(llm_output)

        elif extraction_mode == "Extract to Word":
            try:
                doc_data = parse_word_response(llm_output)

                # Preview in app
                st.markdown(f"## {doc_data.get('title', '')}")
                for section in doc_data.get("sections", []):
                    st.markdown(f"### {section.get('heading', '')}")
                    for block in section.get("content", []):
                        if block.get("type") == "paragraph":
                            st.markdown(block.get("text", ""))
                        elif block.get("type") == "bullets":
                            for item in block.get("items", []):
                                st.markdown(f"- {item}")
                        elif block.get("type") == "table":
                            headers = block.get("headers", [])
                            rows    = block.get("rows", [])
                            if headers and rows:
                                st.dataframe(pd.DataFrame(rows, columns=headers), use_container_width=True)

                # Download
                word_bytes = build_word_doc_from_sections(doc_data)
                st.download_button("⬇️ Download as Word (.docx)",
                                   data=word_bytes,
                                   file_name="document.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Could not parse response: {e}")
                st.code(llm_output)

        elif extraction_mode == "Slide Ready":
            try:
                slides = parse_slide_response(llm_output)
                for i, slide in enumerate(slides):
                    slide["format"] = slide_specs[i]["format"] if i < len(slide_specs) else "Bullet Points"
                render_slides_in_app(slides)
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("⬇️ Download as Word (.docx)",
                                       data=build_word_doc_from_slides(slides),
                                       file_name="slides.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                with col2:
                    plain = "\n\n".join(
                        f"SLIDE {i+1}: {s.get('title','')}\n" + json.dumps(s, indent=2)
                        for i, s in enumerate(slides)
                    )
                    st.download_button("⬇️ Download as .txt",
                                       data=plain.encode("utf-8"),
                                       file_name="slides.txt", mime="text/plain")
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Could not parse response: {e}")
                st.code(llm_output)

        else:  # Summary & Action Items
            try:
                data = parse_summary_response(llm_output)
                st.subheader("📌 Key Takeaways")
                for point in data["summary"]:
                    st.markdown(f"- {point}")
                st.subheader("✅ Action Items")
                if data["action_items"]:
                    action_df = pd.DataFrame(data["action_items"])
                    st.dataframe(action_df, use_container_width=True)
                    buf = io.BytesIO()
                    action_df.to_excel(buf, index=False, engine="openpyxl")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button("⬇️ Download as CSV",
                                           data=action_df.to_csv(index=False).encode("utf-8"),
                                           file_name="action_items.csv", mime="text/csv")
                    with col2:
                        st.download_button("⬇️ Download as Excel",
                                           data=buf.getvalue(),
                                           file_name="action_items.xlsx",
                                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.info("No action items found.")
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Could not parse response: {e}")
                st.code(llm_output)


if __name__ == "__main__":
    main()
